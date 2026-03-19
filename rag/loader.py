from __future__ import annotations

from pathlib import Path

from models.schemas import ResumeDocument, SkippedReason
from utils.exceptions import ParseException


def _read_txt(file_path: Path) -> str:
    return file_path.read_text(encoding="utf-8", errors="ignore").strip()


def _read_pdf(file_path: Path) -> str:
    try:
        from pypdf import PdfReader
    except ImportError as exc:  # pragma: no cover - dependency guard
        raise ParseException("pypdf is required to load PDF files", SkippedReason.PARSE_ERROR_PDF) from exc

    reader = PdfReader(str(file_path))
    text = "\n".join(page.extract_text() or "" for page in reader.pages).strip()
    return text


def _read_docx(file_path: Path) -> str:
    try:
        from docx import Document
    except ImportError as exc:  # pragma: no cover - dependency guard
        raise ParseException("python-docx is required to load DOCX files", SkippedReason.PARSE_ERROR_DOCX) from exc

    doc = Document(str(file_path))
    text = "\n".join(paragraph.text for paragraph in doc.paragraphs).strip()
    return text


def load_resume(file_path: str) -> ResumeDocument:
    path = Path(file_path)
    source_id = path.stem

    if not path.exists():
        raise ParseException(f"file does not exist: {file_path}", SkippedReason.PARSE_ERROR_UNKNOWN)

    extension = path.suffix.lower()
    try:
        if extension == ".txt":
            text = _read_txt(path)
        elif extension == ".pdf":
            text = _read_pdf(path)
        elif extension == ".docx":
            text = _read_docx(path)
        else:
            raise ParseException(
                f"unsupported file extension: {extension or 'unknown'}",
                SkippedReason.PARSE_ERROR_UNKNOWN,
            )
    except ParseException:
        raise
    except Exception as exc:
        reason = {
            ".pdf": SkippedReason.PARSE_ERROR_PDF,
            ".docx": SkippedReason.PARSE_ERROR_DOCX,
            ".txt": SkippedReason.PARSE_ERROR_TXT,
        }.get(extension, SkippedReason.PARSE_ERROR_UNKNOWN)
        raise ParseException(f"failed to parse file: {file_path}", reason) from exc

    if not text:
        reason = {
            ".pdf": SkippedReason.PARSE_ERROR_PDF,
            ".docx": SkippedReason.PARSE_ERROR_DOCX,
            ".txt": SkippedReason.PARSE_ERROR_TXT,
        }.get(extension, SkippedReason.PARSE_ERROR_UNKNOWN)
        raise ParseException(f"empty or unreadable content: {file_path}", reason)

    return ResumeDocument(source_id=source_id, file_path=str(path), text=text)

